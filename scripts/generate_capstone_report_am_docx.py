from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_p(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def main():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    add_title(doc, 'PipeScope․ CI/CD խողովակաշարերի անվտանգության գնահատման գործիք')
    add_subtitle(doc, 'Համալսարանական ավարտական աշխատանքի տեխնիկական հաշվետվություն')
    add_subtitle(doc, 'Տարբերակ՝ 0.1.0 | Տարի՝ 2026')
    doc.add_paragraph()

    add_h1(doc, '1. Գործադիր ամփոփում')
    add_p(
        doc,
        'PipeScope-ը Python-ով իրականացված հրամանային տողի գործիք է, որը նախատեսված է CI/CD միջավայրերում '
        'անվտանգության ռիսկերի հայտնաբերման, դասակարգման և վերլուծական ներկայացման համար։ Գործիքը միավորում է '
        'մի քանի վերլուծական շերտ՝ CI/CD կարգավորումների ստատիկ ստուգում, գաղտնիքների արտահոսքի հայտնաբերում, '
        'Docker/Kubernetes կոնֆիգուրացիոն թերությունների վերլուծություն, ինչպես նաև Jenkins և GitLab հեռավոր '
        'ծառայությունների endpoint-ների քննություն։'
    )
    add_p(
        doc,
        'Արդյունքները տրամադրվում են JSON, PDF և GitHub Actions Summary ձևաչափերով, իսկ AI խորհրդատու ենթահամակարգը '
        'յուրաքանչյուր հայտնաբերման համար գեներացնում է առաջնահերթ և կիրառելի շտկման առաջարկություններ։ Այս ճարտարապետական '
        'մոտեցումը ապահովում է թե՛ մեքենայական ավտոմատացում, թե՛ մարդու կողմից արագ վերլուծելի հաշվետվություն։'
    )

    add_h1(doc, '2. Ներածություն և հիմնավորում')
    add_p(
        doc,
        'Ժամանակակից DevOps միջավայրում CI/CD համակարգերը ստացել են բարձր արտոնություններ և հաճախ հասանելիություն '
        'արտադրական ռեսուրսներին, ինչը դրանք դարձնում է առաջնային հարձակման մակերես։ Աշխատանքային գործընթացների '
        'անվտանգությունը կախված է ոչ միայն ծրագրի աղբյուրային կոդից, այլև pipeline-ների ճշգրիտ և նվազ արտոնություններով '
        'կարգաբերումից։ Փոքր սխալը, օրինակ՝ չսահմանափակված trigger-ը կամ չպինված երրորդ կողմի action-ը, կարող է հանգեցնել '
        'գաղտնիքների արտահոսքի, supply-chain գրոհի կամ չարտոնված տեղակայման։'
    )
    add_p(
        doc,
        'PipeScope նախագծի նպատակն է լրացնել այդ բացը՝ տրամադրելով միասնական հարթակ, որն ավտոմատ կերպով գնահատում է '
        'CI/CD ռիսկերը նույն միջավայրում, որտեղ իրականացվում է կառուցումը և տեղակայումը։'
    )

    add_h1(doc, '3. Նպատակներ և սահմաններ')
    add_h2(doc, '3.1 Հիմնական նպատակներ')
    add_bullets(doc, [
        'CI/CD կարգավորումների անվտանգության սխալների հայտնաբերում (GitHub Actions, GitLab CI, Jenkins):',
        'Jenkins/GitLab ծառայությունների տարբերակային և endpoint մակարդակի ռիսկերի գնահատում:',
        'Գաղտնիքների արտահոսքի բազմաշերտ հայտնաբերում աղբյուրային ֆայլերում և կոնֆիգուրացիաներում:',
        'Dockerfile, docker-compose և Kubernetes մանիֆեստների անվտանգային ստուգում:',
        'NVD API-ի միջոցով CVE-ների կապակցում հայտնաբերված տեխնոլոգիաների հետ:',
        'AI-հիմնված շտկման առաջարկությունների տրամադրում և GitHub Actions Summary ինտեգրում:'
    ])
    add_h2(doc, '3.2 Սահմանափակումներ (v0.1.0)')
    add_bullets(doc, [
        'Դինամիկ DAST թեստավորում ամբողջական ծածկույթով չի իրականացվում:',
        'Կախվածությունների լիարժեք SCA չի փոխարինում մասնագիտացված գործիքներին:',
        'Գործիքը read-only բնույթ ունի և ինքնուրույն չի փոփոխում pipeline ֆայլերը:'
    ])

    add_h1(doc, '4. Համակարգի ճարտարապետություն')
    add_p(
        doc,
        'Համակարգը կառուցված է մոդուլային սկզբունքով։ CLI շերտը ընդունում է մուտքային պարամետրերը և փոխանցում սկանավորման '
        'օրկեստրատորին։ Օրկեստրատորը հերթով կանչում է տեղային և հեռավար սկաներները, միավորում արդյունքները, ապա փոխանցում '
        'դրանք AI enrichment և output ենթահամակարգերին։'
    )
    add_bullets(doc, [
        'CLI շերտ (`pipescope/cli.py`)՝ հրամանների և դրոշների կառավարում:',
        'Օրկեստրատոր (`pipescope/core/scan.py`)՝ մոդուլների կանչ և արդյունքների ագրեգացում:',
        'CI/CD վերլուծիչ (`pipescope/core/cicd_detector.py`)՝ workflow/Jenkinsfile հայտնաբերում և կանոնային գնահատում:',
        'Գաղտնիքների ենթահամակարգ՝ gitleaks + detect-secrets + custom regex + weak-secret ստուգումներ:',
        'Ինֆրաստրուկտուրային սկաներներ՝ Dockerfile, Compose, Kubernetes:',
        'Հեռավար սկաներներ՝ Jenkins/GitLab endpoint reconnaissance և version fingerprinting:',
        'Արդյունքների շերտ՝ JSON, PDF, GitHub Step Summary, Rich terminal UI:'
    ])

    add_h1(doc, '5. Մոդուլային իրականացման վերլուծություն')

    add_h2(doc, '5.1 CLI և թիրախի լուծում')
    add_p(
        doc,
        'CLI-ը իրականացնում է flexible invocation model․ target արգումենտը կարող է բացակայել, եթե տրված են '
        'հեռավար URL-ներ (`--jenkins-url`, `--gitlab-url`, `--web-url`)։ Այս որոշումը թույլ է տալիս remote-only '
        'սցենարներ GitHub Actions միջավայրում առանց repository checkout պարտադիր լինելու։'
    )

    add_h2(doc, '5.2 Սկանավորման օրկեստրացիա')
    add_p(
        doc,
        'Օրկեստրատորը (`scan.py`) ընդունում է `root: Optional[Path]`։ Եթե `root` բացակայում է, տեղային սկաներները '
        'միջանցվում են, իսկ համակարգը անցնում է միայն հեռավար ստուգումների։ Այս մոտեցումը նվազեցնում է runtime սխալների '
        'ռիսկը և ապահովում է հստակ `remote-only` մետատվյալ։'
    )

    add_h2(doc, '5.3 CI/CD դետեկտոր և կանոնների գնահատում')
    add_p(
        doc,
        '`cicd_detector.py` մոդուլը հայտնաբերում է `.github/workflows/*.yml`, `.gitlab-ci.yml`, `Jenkinsfile` '
        'ֆայլերը, ձևավորում դիտարկումների բառարան և գնահատում դրանք TOML կանոնագրքերին համապատասխան։ Գնահատման '
        'օպերատորներն են `contains`, `equals_ci`, `regex`, `any_true`։'
    )

    add_h2(doc, '5.4 GitHub Actions դիտարկումներ')
    add_p(
        doc,
        '`gha_observe.py`-ը նորմալացնում է trigger-ները, permission-ները, self-hosted runner-ի օգտագործումը, '
        'unpinned action հղումները և `curl|bash` pattern-ները։ Այս դիտարկումները ուղղակիորեն կապակցվում են '
        '`cicd_github_actions.toml` կանոններին։'
    )

    add_h2(doc, '5.5 Գաղտնիքների հայտնաբերման բազմաշերտ ենթահամակարգ')
    add_bullets(doc, [
        'Gitleaks adapter՝ subprocess միջոցով արտաքին սկան (եթե հասանելի է):',
        'detect-secrets ինտեգրում՝ entropy և pattern detector-ներով:',
        'Custom TOML regex rules (`secrets.toml`)՝ hardcoded credential-ների համար:',
        'Weak secrets scanner (`weak_secrets.py`)՝ պարզ/կանխադրված արժեքների հայտնաբերում:'
    ])
    add_p(
        doc,
        'Այս բազմաշերտ մոդելը բարձրացնում է հայտնաբերման ամբողջականությունը և թույլ է տալիս graceful degradation '
        'այն դեպքերում, երբ արտաքին գործիքներից որևէ մեկը հասանելի չէ։'
    )

    add_h2(doc, '5.6 Docker և Compose վերլուծիչներ')
    add_p(
        doc,
        '`dockerfile_scanner.py`-ը հայտնաբերում է ռիսկային շաբլոններ՝ unpinned base image, root user, '
        'secrets in ENV, ADD հրամանի ոչ նպատակային կիրառում։ `compose_scanner.py`-ը ստուգում է `privileged: true`, '
        '`network_mode: host`, բացահայտ ports և plaintext secret patterns։'
    )

    add_h2(doc, '5.7 Kubernetes սկանավորում')
    add_p(
        doc,
        '`k8s_scanner.py` մոդուլը ինտեգրվում է Kubescape-ին։ Բինարիի որոնումը կատարվում է երեք փուլով՝ '
        'env override, PATH lookup, կամ fallback `go run` ներկառուցված աղբյուրից։ Սա ապահովում է բարձր տեղափոխելիություն '
        'և նվազ կախվածություն նախնական միջավայրի կոնֆիգուրացիայից։'
    )

    add_h2(doc, '5.8 CVE lookup և cache')
    add_p(
        doc,
        '`cve_lookup.py`-ը օգտագործում է NVD REST API v2 և `lru_cache`՝ կրկնվող հարցումների կրճատման համար։ '
        'Կիրառվում է product name normalization և version fuzzy matching՝ partial version մատչելիության դեպքում '
        'արդյունավետ որոնում ապահովելու համար։'
    )

    add_h2(doc, '5.9 AI խորհրդատու')
    add_p(
        doc,
        '`ai_advisor.py` մոդուլը աշխատում է երկու ռեժիմով՝ (ա) rule-based առաջարկությունների քարտեզ առանց արտաքին API, '
        '(բ) OpenAI key առկայության դեպքում context-aware enrichment։ Յուրաքանչյուր finding ստանում է '
        '`ai_recommendation`, իսկ ընդհանուր արդյունքի համար կառուցվում է executive summary։'
    )

    add_h2(doc, '5.10 GitHub Summary ինտեգրում')
    add_p(
        doc,
        '`github_summary.py` մոդուլը գրում է Markdown հաշվետվություն `$GITHUB_STEP_SUMMARY` միջավայրային փոփոխականի '
        'ֆայլում՝ severity breakdown, findings table, AI summary և CVE հղումներ։ Այսպիսով արդյունքները հասանելի են '
        'անմիջապես Actions UI-ում՝ առանց հավելյալ արտեֆակտների բացման։'
    )

    add_h1(doc, '6. Կանոնների (TOML) համակարգ')
    add_p(
        doc,
        'PipeScope-ում կանոնները առանձնացված են կոդից և պահվում են `pipescope/rules/` պանակում։ Այս մոտեցումը '
        'թույլ է տալիս անվտանգային գիտելիքի թարմացումներ կատարել առանց core business logic փոփոխելու։'
    )
    add_bullets(doc, [
        '`cicd_github_actions.toml`՝ trigger/permission/action pinning/run pattern կանոններ:',
        '`cicd_gitlab.toml`՝ GitLab pipeline անվտանգության կանոններ:',
        '`cicd_jenkins.toml`՝ Jenkinsfile անվտանգության ստուգումներ:',
        '`secrets.toml`՝ regex-հիմնված hardcoded credential կանոններ:'
    ])

    add_h1(doc, '7. GitHub Actions ինտեգրման ճարտարապետություն')
    add_p(
        doc,
        'Նախագծում առկա է reusable composite action (`.github/actions/pipescope/action.yml`) և օրինակ workflow '
        '(`.github/workflows/pipescope.yml`)։ Գլխավոր job-երը ներառում են repository scan, optional jenkins scan '
        'և consolidated summary։ Action-ը վերադարձնում է count-based outputs՝ findings, high, critical։'
    )

    add_h1(doc, '8. Տեխնոլոգիական փաթեթ և կախվածություններ')
    add_bullets(doc, [
        'Python 3.10+ (գործնականում՝ venv Python 3.13.7):',
        'Typer՝ CLI API:',
        'PyYAML՝ YAML parsing:',
        'Rich՝ terminal UI:',
        'detect-secrets՝ secret scanners:',
        'requests՝ HTTP և NVD API client:',
        'reportlab՝ PDF գեներացում:',
        'python-docx՝ DOCX հաշվետվությունների գեներացում:'
    ])

    add_h1(doc, '9. Թեստավորման ռազմավարություն')
    add_p(
        doc,
        'Թեստավորումը իրականացվել է երեք շերտով՝ (1) մոդուլային/սմոք թեստեր (`tests/`), '
        '(2) ինտեգրացիոն փորձարկումներ փորձնական կոնֆիգուրացիաներով, (3) CI վավերացում GitHub Actions-ում։ '
        'Հատուկ ստուգվել են AI enrichment-ի no-key fallback վարքը, GitHub Summary writer-ի fail-safe մեխանիզմը '
        'և rule-based evaluator-ի կայունությունը։'
    )

    add_h1(doc, '10. Փաթեթավորում և բաշխում')
    add_p(
        doc,
        'Նախագիծը կառուցված է `pyproject.toml`-ով (`setuptools` + `wheel`)։ Իրականացման ընթացքում հայտնաբերվեց '
        'կարևոր packaging խնդիր՝ wheel-ում non-Python asset-ների (`rules/*.toml`, `wordlists/*.txt`) բացակայություն, '
        'որը շտկվեց package-data կոնֆիգուրացիայով։ Այս փոփոխությունը կրիտիկական էր GitHub Actions runtime '
        'սխալների վերացման համար։'
    )

    add_h1(doc, '11. Հիմնական տեխնիկական որոշումներ')
    add_bullets(doc, [
        'Optional target պարամետր՝ remote-only սկանավորման պաշտոնական աջակցությամբ:',
        'Rules-as-data մոտեցում՝ TOML կանոնագրքերով ընդլայնելիություն:',
        'Multi-layer secret detection՝ recall-ի բարձրացման նպատակով:',
        'LRU cache CVE lookup-ում՝ API ծանրաբեռնվածության նվազեցում:',
        'Kubescape fallback (go run)՝ գործարկելիություն heterogeneous միջավայրերում:'
    ])

    add_h1(doc, '12. Սահմանափակումներ և հետագա աշխատանք')
    add_bullets(doc, [
        'SARIF արտածման բացակայություն (պլանավորվում է Code Scanning ինտեգրման համար):',
        'Bitbucket Pipelines-ի բացակայող աջակցություն:',
        'NVD rate-limit կախվածություն առանց API key-ի:',
        'Git history secret scan-ի սահմանափակ ծածկույթ current mode-ում:',
        'Կեղծ դրականների օպտիմալացման լրացուցիչ անհրաժեշտություն weak-secret հուրիստիկայում:'
    ])
    add_p(
        doc,
        'Հետագա փուլում նպատակահարմար է ավելացնել SARIF, HTML interactive report, plugin interface և '
        'ավելի խորացված policy-as-code շերտ՝ կազմակերպական անվտանգության չափանիշներին համապատասխանեցման համար։'
    )

    add_h1(doc, '13. Եզրակացություն')
    add_p(
        doc,
        'PipeScope ավարտական նախագիծը ցուցադրում է DevSecOps անվտանգության ավտոմատացման ամբողջական մոտեցում, '
        'որտեղ մեկ գործիքի շրջանակում համադրվում են pipeline կոնֆիգուրացիայի վերլուծությունը, գաղտնիքների հայտնաբերումը, '
        'ինֆրաստրուկտուրային ստուգումները, CVE enrichment-ը և AI-հիմնված շտկման առաջարկությունները։ '
        'Ճարտարապետական լուծումները մոդուլային են, ընդլայնելի և կիրառելի ինչպես ուսումնական, այնպես էլ արտադրական միջավայրերում։'
    )

    add_h1(doc, '14. Օգտագործված աղբյուրներ (ընտրովի ցուցակ)')
    add_bullets(doc, [
        'NIST NVD API v2 փաստաթղթեր:',
        'GitHub Actions Security Hardening ուղեցույց:',
        'MITRE ATT&CK և CWE դասակարգումներ:',
        'Gitleaks, detect-secrets, Kubescape պաշտոնական փաստաթղթեր:',
        'Python Packaging User Guide և setuptools փաստաթղթեր:'
    ])

    output_path = 'CAPSTONE_REPORT_AM.docx'
    doc.save(output_path)
    print(f'Created: {output_path}')


if __name__ == '__main__':
    main()
